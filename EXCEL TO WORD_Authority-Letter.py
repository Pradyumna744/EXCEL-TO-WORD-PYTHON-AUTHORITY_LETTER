import os
import pandas as pd
import datetime
from docx import Document
from tkinter import Tk, Text, filedialog, Scrollbar
from tkinter.ttk import Progressbar, Label, Button, Entry, Frame
import ttkbootstrap as ttk
from ttkbootstrap.dialogs import Messagebox


def sanitize_filename(filename):
    """Remove or replace invalid characters in filenames."""
    invalid_chars = r'<>:"/\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, "")
    return filename.strip()


def mail_merge(template_path, excel_path, output_folder, log_text, progress_bar):
    try:
        # Load the Excel file, treating all data as strings to handle mixed formats
        df = pd.read_excel(excel_path, dtype=str)
        df.columns = df.columns.str.strip()  # Remove extra spaces from column names
        df.fillna("", inplace=True)  # Replace NaN values with empty strings

        # Convert all datetime columns to dd-mm-yyyy format, except DOA and DOD
        for col in df.columns:
            try:
                # Skip columns that are DOA or DOD and format them separately
                if col == "DOA" or col == "DOD":
                    df[col] = pd.to_datetime(df[col], errors='coerce').dt.strftime('%d-%m-%Y')
                    continue
                
                # Attempt to convert column to datetime, if possible
                if pd.to_datetime(df[col], errors='coerce').notna().all():
                    df[col] = pd.to_datetime(df[col], errors='coerce').dt.strftime('%d-%m-%Y')
            except Exception as e:
                pass  # Ignore columns that aren't datetime-related

        total_records = len(df)
        log_text.insert("end", f"Loaded {total_records} records from the Excel file.\n")
        progress_bar["maximum"] = total_records

        # Get the current date in dd-mm-yyyy format
        current_date = datetime.datetime.now().strftime('%d-%m-%Y')
        df["Current Date"] = current_date

        # Process each row in the Excel file
        for index, row in df.iterrows():
            try:
                case_assign = sanitize_filename(str(row.get("Case Assign", "Case Assign")).strip())
                claim_no = sanitize_filename(str(row.get("ClNo", "")).strip())
                claim_id = sanitize_filename(str(row.get("ClId", "")).strip())
                policy_no = sanitize_filename(str(row.get("PN", "")).strip())
                insured_name = sanitize_filename(str(row.get("Insured Name", "")).strip())

                # Determine folder identifier
                folder_identifier = (
                    f"{claim_no} {insured_name}" if claim_no else
                    f"{claim_id} {insured_name}" if claim_id else
                    f"{policy_no} {insured_name}" if policy_no else
                    insured_name
                ).strip()

                if not folder_identifier:
                    log_text.insert("end", f"Skipping record {index+1}: Missing identifiers.\n")
                    continue

                # Create folders grouped by "Case Assign" and folder identifier
                case_folder_path = os.path.join(output_folder, case_assign, folder_identifier)
                os.makedirs(case_folder_path, exist_ok=True)

                # Load and process Word template
                doc = Document(template_path)
                for paragraph in doc.paragraphs:
                    for key in df.columns:
                        placeholder = f"[{key}]"
                        value = str(row.get(key, "")).strip()
                        if value in {"", "###", None}:
                            value = ""
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

                for table in doc.tables:
                    for row_cells in table.rows:
                        for cell in row_cells.cells:
                            for key in df.columns:
                                placeholder = f"[{key}]"
                                value = str(row.get(key, "")).strip()
                                if value in {"", "###", None}:
                                    value = ""
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, value)

                # Replace the current date placeholder
                for paragraph in doc.paragraphs:
                    if "[Current Date]" in paragraph.text:
                        paragraph.text = paragraph.text.replace("[Current Date]", current_date)

                for table in doc.tables:
                    for row_cells in table.rows:
                        for cell in row_cells.cells:
                            if "[Current Date]" in cell.text:
                                cell.text = cell.text.replace("[Current Date]", current_date)

                # Save the output document inside the corresponding folder
                output_file = os.path.join(case_folder_path, "Authority Letter.docx")
                doc.save(output_file)

                # Update log and progress bar
                log_text.insert("end", f"Processed: {folder_identifier} under {case_assign}\n")
                progress_bar["value"] = index + 1
                log_text.see("end")

            except Exception as row_error:
                log_text.insert("end", f"Error processing record {index+1}: {row_error}\n")
                continue

        log_text.insert("end", f"Mail merge completed! Files saved in: {output_folder}\n")
        Messagebox.show_info("Success", "Mail merge completed successfully!")

    except Exception as e:
        log_text.insert("end", f"Critical Error: {e}\n")
        Messagebox.show_error("Error", f"An error occurred: {e}")


# Tkinter GUI Setup
def select_template():
    """Open file dialog to select template"""
    template_path.set(filedialog.askopenfilename(filetypes=[("Word files", "*.docx")]))

    log_text.insert("end", f"Template selected: {template_path.get()}\n")
    log_text.see("end")


def select_excel():
    """Open file dialog to select Excel file"""
    excel_path.set(filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")]))

    log_text.insert("end", f"Excel file selected: {excel_path.get()}\n")
    log_text.see("end")


def select_output_folder():
    """Open folder dialog to select output folder"""
    folder = filedialog.askdirectory()
    if folder:
        output_folder.set(folder)
        log_text.insert("end", f"Output folder selected: {folder}\n")
        log_text.see("end")


def start_mail_merge():
    """Start the mail merge operation"""
    if not template_path.get() or not excel_path.get() or not output_folder.get():
        Messagebox.show_error("Error", "Please select all files and the output folder.")
        return

    progress_bar["value"] = 0
    log_text.delete("1.0", "end")
    mail_merge(template_path.get(), excel_path.get(), output_folder.get(), log_text, progress_bar)


# Initialize the Tkinter window
root = Tk()
root.title("Mail Merge Application")

# Define variables
template_path = ttk.StringVar()
excel_path = ttk.StringVar()
output_folder = ttk.StringVar()

# Create GUI components
frame = Frame(root)
frame.pack(padx=10, pady=10)

Label(frame, text="Template File:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
Entry(frame, textvariable=template_path, width=40).grid(row=0, column=1, padx=5, pady=5)
Button(frame, text="Browse", command=select_template).grid(row=0, column=2, padx=5, pady=5)

Label(frame, text="Excel File:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
Entry(frame, textvariable=excel_path, width=40).grid(row=1, column=1, padx=5, pady=5)
Button(frame, text="Browse", command=select_excel).grid(row=1, column=2, padx=5, pady=5)

Label(frame, text="Output Folder:").grid(row=2, column=0, sticky="e", padx=5, pady=5)
Entry(frame, textvariable=output_folder, width=40).grid(row=2, column=1, padx=5, pady=5)
Button(frame, text="Browse", command=select_output_folder).grid(row=2, column=2, padx=5, pady=5)

Button(root, text="Start Mail Merge", command=start_mail_merge).pack(padx=10, pady=10)

# Log and progress bar
log_text = Text(root, height=15, width=80)
log_text.pack(padx=10, pady=5)

progress_bar = Progressbar(root, length=100, mode="determinate")
progress_bar.pack(padx=10, pady=5)

root.mainloop()
